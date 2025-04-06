"""
# Excel to Word 変換ツール

このスクリプトは、Excelファイルからデータを読み込み、各人の情報をもとにWordドキュメントを生成します。
ゴルフ場利用税の非課税申請書を自動生成するためのツールです。

## 機能
- Excelファイルから氏名、住所、生年月日のデータを読み込む
- 各人ごとにWordドキュメントを生成する
- 生成したドキュメントに適切なフォーマットを適用する

## 使い方
1. 設定セクションで入力ファイルパス、出力ディレクトリ、シート名を設定する
2. スクリプトを実行する
3. 指定した出力ディレクトリに各人のWordドキュメントが生成される

## 必要なライブラリ
- pandas: データ処理用
- openpyxl: Excelファイル操作用
- python-docx: Wordドキュメント操作用
- datetime: 日付処理用
"""

import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# python-dotenvがインストールされている場合は.envファイルから環境変数を読み込む
try:
    from dotenv import load_dotenv
    load_dotenv()
    print(".envファイルから環境変数を読み込みました")
except ImportError:
    print("python-dotenvがインストールされていないため、.envファイルは読み込まれません")
    print("pip install python-dotenvでインストールできます")

# 設定セクション（環境変数から取得）
INPUT_FILE_PATH = os.environ.get('INPUT_FILE_PATH', 'input.xlsx')  # 入力Excelファイルのパス
OUTPUT_DIRECTORY = os.environ.get('OUTPUT_DIRECTORY', 'output')    # 出力ディレクトリ
SHEET_NAME = os.environ.get('SHEET_NAME', '全部員')                # データが含まれるシート名
GOVERNOR_NAME = os.environ.get('GOVERNOR_NAME', '知事名')          # 知事の名前
GOLF_COURSE_NAME = os.environ.get('GOLF_COURSE_NAME', '')          # ゴルフ場名
USAGE_DATE = os.environ.get('USAGE_DATE', '年　　　月　　　日')    # 利用年月日

def setup_environment():
    """出力ディレクトリが存在しない場合は作成します。"""
    try:
        if not os.path.exists(OUTPUT_DIRECTORY):
            os.makedirs(OUTPUT_DIRECTORY)
            print("出力ディレクトリを作成しました")
        return True
    except Exception as e:
        print("エラー: 出力ディレクトリの作成に失敗しました")
        return False

def load_excel_data(file_path, sheet_name):
    """Excelファイルからデータを読み込みます。"""
    try:
        print("Excelファイルからデータを読み込んでいます...")
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        print("データの読み込みが完了しました")
        return df
    except Exception as e:
        print("エラー: データの読み込み中に問題が発生しました")
        return None

def extract_unique_data(df):
    """データフレームから氏名、住所、生年月日のユニークな値を抽出します。"""
    try:
        names = df['氏名'].unique()
        addresses = df['住所'].unique()
        birthdays = df['生年月日'].unique()
        return names, addresses, birthdays
    except Exception as e:
        print("エラー: データの抽出中に問題が発生しました")
        return [], [], []

def parse_date(date_string):
    """日付文字列をパースして年、月、日に分解します。"""
    try:
        date = datetime.strptime(str(date_string), '%Y-%m-%d %H:%M:%S')
        return date.year, date.month, date.day
    except ValueError:
        try:
            date = datetime.strptime(str(date_string), '%Y-%m-%d')
            return date.year, date.month, date.day
        except Exception:
            print("エラー: 日付のパースに失敗しました")
            return None, None, None
    except Exception:
        print("エラー: 日付の処理中に問題が発生しました")
        return None, None, None

def set_cell_borders_black(cell):
    """セルの枠線を黒色に設定します。"""
    tcPr = cell._element.get_or_add_tcPr()

    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:color'), '000000')
    tcPr.append(top)

    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '4')
    left.set(qn('w:color'), '000000')
    tcPr.append(left)

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), '000000')
    tcPr.append(bottom)

    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'single')
    right.set(qn('w:sz'), '4')
    right.set(qn('w:color'), '000000')
    tcPr.append(right)

def apply_font_settings(file_path):
    """ドキュメントのフォント設定を適用します。"""
    try:
        doc = Document(file_path)

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.name = 'ＭＳ 明朝'

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.name = 'ＭＳ 明朝'

        doc.save(file_path)
        print("フォント設定を適用しました")
    except Exception:
        print("エラー: フォント設定の適用中に問題が発生しました")

def create_word_document(name, address, birthday, index, governor_name):
    """Wordドキュメントを作成します。"""
    try:
        doc = Document()

        doc.add_paragraph('第38号の5様式\n\n')
        doc.add_paragraph('ゴルフ場利用税非課税申請書\n').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f'　　石川県知事　{governor_name}様\n')
        doc.add_paragraph('　私は、ゴルフ場利用税非課税対象者に該当しますので、申請します。')

        table = doc.add_table(rows=9, cols=2)

        for row in table.rows:
            for cell in row.cells:
                set_cell_borders_black(cell)

        for row in table.rows:
            row.cells[0].width = Inches(1.4)
            row.cells[1].width = Inches(4.3)

        year, month, day = parse_date(birthday)
        if year is None:
            birthday_str = ""
        else:
            japanese_year = year - 1988
            birthday_str = f'大・昭・平{japanese_year}年{month}月{day}日生'

        rows_contents = [
            ('利用ゴルフ場名', GOLF_COURSE_NAME),
            ('利用年月日', USAGE_DATE),
            ('住所', address),
            ('区分', '□　メンバー　　□　ビジター'),
            ('氏名', name),
            ('生年月日', birthday_str),
            ('非課税等適用区分', '□70歳以上　□18歳未満　□障害者等\n☑︎教育活動等　□国民スポーツ大会　□スポーツマスターズ等'),
            ('証明書の種類', '□運転免許証　☑︎学生証　□職員証　□パスポート\n□障害者手帳等　□学校長の証明　□教育委員会の証明\n□その他(　　　　　　　　)'),
            ('備考','')
        ]

        for row, (label, content) in zip(table.rows, rows_contents):
            row.cells[0].text = label
            row.cells[1].text = content
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        remarks = '''備考\n
        1　該当する□の中にレ点を付けてください。
　　　   2　70歳以上、18歳未満及び障害者等の方は、この申請書を、利用するゴルフ場が最初の
　　　利用である場合にゴルフ場に提出してください。また、受付の際に非課税利用に該当する
　　　ことを証明する証明書をゴルフ場に提示してください。
　　　3　教育活動等、国民スポーツ大会、スポーツマスターズ等の利用の場合は、利用の都度
　　　この申請書を提出してください。その際には、受付に非課税・課税免除利用に該当するこ
　　　とを証明する証明書をゴルフ場に提出してください。
　　　4　この申請書を提出しない場合、2又は3の証明書を提示又は提出しない場合は、非課
　　　税・課税免除の適用を受けられない場合があります。
'''
        doc.add_paragraph(remarks).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        file_name = f"{index+1:03d}_{name}.docx"
        file_path = os.path.join(OUTPUT_DIRECTORY, file_name)
        doc.save(file_path)

        apply_font_settings(file_path)

        print("ドキュメントを作成しました")
        return file_path
    except Exception:
        print("エラー: ドキュメントの作成中に問題が発生しました")
        return None

def main():
    """メイン処理を実行します。"""
    print("Excel to Word 変換ツールを開始します")

    # 環境のセットアップ
    if not setup_environment():
        print("環境のセットアップに失敗しました。処理を中止します。")
        return

    # Excelデータの読み込み
    df = load_excel_data(INPUT_FILE_PATH, SHEET_NAME)
    if df is None:
        print("Excelデータの読み込みに失敗しました。処理を中止します。")
        return

    # ユニークなデータの抽出
    names, addresses, birthdays = extract_unique_data(df)
    if len(names) == 0:
        print("データの抽出に失敗しました。処理を中止します。")
        return

    # 各人のWordドキュメントを作成
    success_count = 0
    for i in range(len(names)):
        file_path = create_word_document(names[i], addresses[i], birthdays[i], i, GOVERNOR_NAME)
        if file_path:
            success_count += 1

    print(f"処理が完了しました。{success_count}件のドキュメントを作成しました。")

if __name__ == "__main__":
    main()
